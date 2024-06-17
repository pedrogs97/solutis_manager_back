"""Task"""

import base64
import math
import os
from datetime import date
from os.path import dirname

import jinja2
import pdfkit
from invoke import task
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from sqlalchemy import Table

from src.config import CONTRACT_UPLOAD_TEST_DIR, DEFAULT_DATE_FORMAT, TEMPLATE_DIR
from src.database import Base, Engine, Session_db
from src.report.service import ReportService
from src.utils import get_file_paths, read_file


@task
def makemigrations(cmd, message):
    """Generate alembic migrations"""
    cmd.run(f'alembic revision --autogenerate -m "{message}"')


@task
def migrate(cmd):
    """Execute alembic migrations"""
    cmd.run("alembic upgrade head")


@task
def revertmigration(cmd):
    """Revert last alembic migration"""
    cmd.run("alembic downgrade -1")


@task
def run(cmd):
    """Run application"""
    cmd.run("uvicorn src.main:appAPI --port 8000 --reload")


@task
def loaddata(cmd, module: str, fixture: str = None):
    """Load fixtures to database"""
    fixtures_directory: str = dirname(__file__) + f"/src/{module}/fixtures/"
    try:
        if fixture:
            file_path: str = fixtures_directory + fixture + ".json"
            fixtures: list = read_file(file_path=file_path)
            for item in fixtures:
                table_name = item.pop("table", "")
                if table_name == "":
                    cmd.run("echo table not found")
                    return
                table = Table(table_name, Base.metadata, autoload_with=Engine)
                with Session_db() as session:
                    if not session.execute(
                        table.select().where(table.c.id == item["id"])
                    ).first():
                        session.execute(table.insert().values(item))
                        session.commit()
            cmd.run(f"echo {len(fixtures)} fixtures applied on {table_name}")
        else:
            file_paths: list = get_file_paths(fixtures_directory)
            for file_path in file_paths:
                fixtures: list = read_file(file_path=fixtures_directory + file_path)
                for item in fixtures:
                    table_name = item.pop("table", "")
                    if table_name == "":
                        cmd.run("echo table not found")
                        return
                    table = Table(table_name, Base.metadata, autoload_with=Engine)
                    with Session_db() as session:
                        if not session.execute(
                            table.select().where(table.c.id == item["id"])
                        ).first():
                            session.execute(table.insert().values(item))
                            session.commit()
                cmd.run(f"echo {len(fixtures)} fixtures applied on {table_name}")
    except FileNotFoundError:
        cmd.run(f"echo fixture {fixture} not found")


def __contract():
    """
    Contract
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "comodato.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        marital_status="CASADO",
        cc="23412-1",
        manager="Hericles Bitencurt",
        business_executive="Janaina Bitencurt",
        project="Solutis",
        workload="Home Office",
        register_number="sa321431",
        serial_number="2151232",
        description="Macbook Air",
        accessories="N/A",
        ms_office="Sim",
        pattern="Studio",
        operational_system="MacOS",
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        witnesses=[
            {
                "full_name": "Testemunha 1 teste",
                "taxpayer_identification": "000.000.000-00",
            },
            {
                "full_name": "Testemunha 2 teste",
                "taxpayer_identification": "111.111.111-11",
            },
        ],
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "contract_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


def __termination():
    """
    Termination
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "distrato_comodato.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        marital_status="CASADO",
        cc="23412-1",
        manager="Hericles Bitencurt",
        business_executive="Janaina Bitencurt",
        project="Solutis",
        workload="Home Office",
        register_number="sa321431",
        serial_number="2151232",
        description="Macbook Air",
        accessories="N/A",
        ms_office="Sim",
        pattern="Studio",
        operational_system="MacOS",
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        witnesses=[
            {
                "full_name": "Testemunha 1 teste",
                "taxpayer_identification": "000.000.000-00",
            },
            {
                "full_name": "Testemunha 2 teste",
                "taxpayer_identification": "111.111.111-11",
            },
        ],
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "termination_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


def __termination_pj():
    """
    Termination PJ
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "distrato_comodato_pj.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        marital_status="CASADO",
        cc="23412-1",
        manager="Hericles Bitencurt",
        business_executive="Janaina Bitencurt",
        project="Solutis",
        workload="Home Office",
        register_number="sa321431",
        serial_number="2151232",
        description="Macbook Air",
        accessories="N/A",
        ms_office="Sim",
        pattern="Studio",
        operational_system="MacOS",
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        witnesses=[
            {
                "full_name": "Testemunha 1 teste",
                "taxpayer_identification": "000.000.000-00",
            },
            {
                "full_name": "Testemunha 2 teste",
                "taxpayer_identification": "111.111.111-11",
            },
        ],
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "termination_pj_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


def __contract_pj():
    """
    Contract
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "comodato_pj.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        company="MEI teste",
        cnpj="13.185.790/0001-79",
        company_address="Rua da esquina, 31, Pituba, Salvador",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        marital_status="CASADO",
        cc="23412-1",
        manager="Hericles Bitencurt",
        business_executive="Janaina Bitencurt",
        project="Solutis",
        workload="Home Office",
        contract_date=date.today().strftime(DEFAULT_DATE_FORMAT),
        goal="objetivo teste",
        register_number="sa321431",
        serial_number="2151232",
        description="Macbook Air",
        accessories="N/A",
        ms_office="Sim",
        pattern="Studio",
        operational_system="MacOS",
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        witnesses=[
            {
                "full_name": "Testemunha 1 teste",
                "taxpayer_identification": "000.000.000-00",
            },
            {
                "full_name": "Testemunha 2 teste",
                "taxpayer_identification": "111.111.111-11",
            },
        ],
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_pj_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_pj_test.html"),
        encoding="utf-8",
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "contract_pj_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_pj_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


def __term():
    """
    Term
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "termo.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        cc="23412-1",
        manager="Hericles Bitencurt",
        project="Solutis",
        description="Camisa",
        size="N/A",
        quantity=1,
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "term_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


def __termination_term():
    """
    Termination term
    """
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "distrato_termo.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    signed_file = open("./src/static/images/signed.png", "rb")
    signed_encoded_string = (
        str(base64.b64encode(signed_file.read())).replace("b'", "").replace("'", "")
    )
    date_image = open("./src/static/images/date.jpeg", "rb")
    date_image_encoded_string = (
        str(base64.b64encode(date_image.read())).replace("b'", "").replace("'", "")
    )
    n_glpi_file = open("./src/static/images/n_glpi.png", "rb")
    n_glpi_encoded_string = (
        str(base64.b64encode(n_glpi_file.read())).replace("b'", "").replace("'", "")
    )
    n_termo_file = open("./src/static/images/n_termo.png", "rb")
    n_termo_encoded_string = (
        str(base64.b64encode(n_termo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        number="35161343241",
        glpi_number="GLPI 41325",
        full_name="Abmael da Silva",
        taxpayer_identification="222.222.222-22",
        national_identification="22222222-22",
        address="Rua da esquina, 31, Alpha Vile, Salvador",
        nationality="BRASILEIRO",
        role="Desenvolvedor",
        cc="23412-1",
        manager="Hericles Bitencurt",
        project="Solutis",
        description="Camisa",
        size="N/A",
        quantity=1,
        value="46.000,00",
        date=date.today().strftime(DEFAULT_DATE_FORMAT),
        signed=f"data:image/png;base64,{signed_encoded_string}",
        date_image=f"data:image/png;base64,{date_image_encoded_string}",
        n_glpi=f"data:image/png;base64,{n_glpi_encoded_string}",
        n_termo=f"data:image/png;base64,{n_termo_encoded_string}",
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        location="Salvador - BA",
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "termination_term_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    signed_file.close()
    date_image.close()
    n_glpi_file.close()
    n_termo_file.close()


@task
def test(cmd):
    """
    Convert html to pdf using pdfkit which is a wrapper of wkhtmltopdf
    """
    cmd.run("echo conversion started")
    __term()
    __termination_term()
    __contract()
    __termination()
    __contract_pj()
    __termination_pj()
    cmd.run("echo conversion finished")


@task
def testverification(cmd):
    """
    Convert html to pdf using pdfkit which is a wrapper of wkhtmltopdf
    """
    cmd.run("echo conversion started")
    template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATE_DIR)
    template_env = jinja2.Environment(loader=template_loader)
    template_file = "verification.html"
    template = template_env.get_template(template_file)
    logo_file = open("./src/static/images/ri_1.png", "rb")
    logo_encoded_string = (
        str(base64.b64encode(logo_file.read())).replace("b'", "").replace("'", "")
    )
    output_text = template.render(
        ri_1=f"data:image/png;base64,{logo_encoded_string}",
        verifications=[
            {
                "question": "Carregador - Estado do conector",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 42,
                    },
                    {
                        "checked": False,
                        "option": "Danificado",
                        "id": 43,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 44,
                    },
                ],
            },
            {
                "question": "Serial do carregador",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 45,
                    },
                    {
                        "checked": False,
                        "option": "Danificado",
                        "id": 46,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 47,
                    },
                ],
            },
            {
                "question": "Carregador - Estado do cabo de energia",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 48,
                    },
                    {
                        "checked": False,
                        "option": "Danificado",
                        "id": 49,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 50,
                    },
                ],
            },
            {
                "question": "Borrachas de proteção",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 1,
                    },
                    {
                        "checked": False,
                        "option": "Danificado",
                        "id": 2,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 3,
                    },
                ],
            },
            {
                "question": "Estado do LCD",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 4,
                    },
                    {
                        "checked": False,
                        "option": "Riscado",
                        "id": 5,
                    },
                    {
                        "checked": False,
                        "option": "Quebrado",
                        "id": 6,
                    },
                    {
                        "checked": False,
                        "option": "Não liga",
                        "id": 7,
                    },
                ],
            },
            {
                "question": "Estado do touchpad",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 8,
                    },
                    {
                        "checked": False,
                        "option": "Riscado",
                        "id": 9,
                    },
                    {
                        "checked": False,
                        "option": "Não funciona",
                        "id": 10,
                    },
                    {
                        "checked": False,
                        "option": "Impossível verificar pois não liga",
                        "id": 11,
                    },
                ],
            },
            {
                "question": "Estado da carcaça",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 12,
                    },
                    {
                        "checked": False,
                        "option": "Riscado",
                        "id": 13,
                    },
                    {
                        "checked": False,
                        "option": "Quebrado",
                        "id": 14,
                    },
                ],
            },
            {
                "question": "Estado do Teclado",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 15,
                    },
                    {
                        "checked": False,
                        "option": "Faltando tecla",
                        "id": 16,
                    },
                    {
                        "checked": False,
                        "option": "Não funciona",
                        "id": 17,
                    },
                ],
            },
            {
                "question": "Carregador Bateria (entrada)",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 18,
                    },
                    {
                        "checked": False,
                        "option": "Danificado",
                        "id": 19,
                    },
                ],
            },
            {
                "question": "Conectores USB/Ethernet/Modem",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 20,
                    },
                    {
                        "checked": False,
                        "option": "Danificados",
                        "id": 21,
                    },
                ],
            },
            {
                "question": "Drive de Disco CD - DVD - DVDRW",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 22,
                    },
                    {
                        "checked": False,
                        "option": "Não funciona",
                        "id": 23,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 24,
                    },
                ],
            },
            {
                "question": "Estado da carcaça (atrás da tela)",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 25,
                    },
                    {
                        "checked": False,
                        "option": "Riscado",
                        "id": 26,
                    },
                    {
                        "checked": False,
                        "option": "Quebrado",
                        "id": 27,
                    },
                ],
            },
            {
                "question": "Bateria - Verificar serial (no notebook)",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 28,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 29,
                    },
                ],
            },
            {
                "question": "Borracha de apoio",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 30,
                    },
                    {
                        "checked": False,
                        "option": "Faltam",
                        "id": 31,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 32,
                    },
                ],
            },
            {
                "question": "Serial da máquina",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 33,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 34,
                    },
                ],
            },
            {
                "question": "Parafusos da carcaça",
                "options": [
                    {
                        "checked": True,
                        "option": "Ok",
                        "id": 35,
                    },
                    {
                        "checked": False,
                        "option": "Faltam",
                        "id": 36,
                    },
                    {
                        "checked": False,
                        "option": "Não tem",
                        "id": 37,
                    },
                ],
            },
        ],
    )

    if not os.path.exists(CONTRACT_UPLOAD_TEST_DIR):
        os.mkdir(CONTRACT_UPLOAD_TEST_DIR)

    html_path = os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(output_text)

    options = {
        "page-size": "A4",
        "enable-local-file-access": None,
    }

    with open(
        os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"), encoding="utf-8"
    ) as f:
        pdfkit.from_file(
            f,
            os.path.join(CONTRACT_UPLOAD_TEST_DIR, "verification_test.pdf"),
            options=options,
        )
    os.remove(os.path.join(CONTRACT_UPLOAD_TEST_DIR, "template_test.html"))
    logo_file.close()
    cmd.run("echo conversion finished")


@task
def testreport(cmd):
    """
    Convert html to pdf using pdfkit which is a wrapper of wkhtmltopdf
    """
    cmd.run("echo creation started")
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches
    from PIL import Image, ImageDraw, ImageFont

    # Função para criar uma imagem com uma forma geométrica
    def create_shape_image_with_text(shape, filename, text, size=(150, 150)):
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default(size=14)  # Usar a fonte padrão do Pillow

        # Calcular o tamanho do texto
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]

        if shape == "circle":
            draw.ellipse([10, 10, size[0] - 10, size[1] - 10], outline="black", width=2)
            draw.multiline_text(
                ((size[0] - text_width) / 2, (size[1] - text_height) / 2),
                text,
                fill="black",
                font=font,
                align="center",
            )
        elif shape == "rectangle":
            draw.rectangle(
                [10, 10, size[0] - 10, size[1] - 10], outline="black", width=2
            )
            draw.multiline_text(
                ((size[0] - text_width) / 2, (size[1] - text_height) / 2),
                text,
                fill="black",
                font=font,
                align="center",
            )
        elif shape == "diamond":
            draw.polygon(
                [
                    (size[0] // 2, 10),
                    (size[0] - 10, size[1] // 2),
                    (size[0] // 2, size[1] - 10),
                    (10, size[1] // 2),
                ],
                outline="black",
                width=2,
            )
            draw.multiline_text(
                ((size[0] - text_width) / 2, (size[1] - text_height) / 2),
                text,
                fill="black",
                font=font,
                align="center",
            )
        elif shape == "hexagon":
            draw.polygon(
                [
                    (size[0] // 2, 10),
                    (size[0] - 10, size[1] // 4),
                    (size[0] - 10, 3 * size[1] // 4),
                    (size[0] // 2, size[1] - 10),
                    (10, 3 * size[1] // 4),
                    (10, size[1] // 4),
                ],
                outline="black",
                width=2,
            )
            draw.multiline_text(
                ((size[0] - text_width) / 2, (size[1] - text_height) / 2),
                text,
                fill="black",
                font=font,
                align="center",
            )
        elif shape == "arrow":
            draw.polygon(
                [
                    (size[0] // 4, size[1] // 2),
                    (3 * size[0] // 4, size[1] // 2),
                    (3 * size[0] // 4, size[1] // 4),
                    (size[0] - 10, size[1] // 2),
                    (3 * size[0] // 4, 3 * size[1] // 4),
                    (3 * size[0] // 4, size[1] // 2),
                ],
                outline="black",
                width=2,
            )

        img.save(filename)

    # Texto a ser inserido dentro das figuras
    text = "Comodato teste 2024/06/12"

    # Criar imagens das formas geométricas com texto
    create_shape_image_with_text("circle", "circle.png", text)
    create_shape_image_with_text("rectangle", "rectangle.png", text)
    create_shape_image_with_text("diamond", "diamond.png", text)
    create_shape_image_with_text("hexagon", "hexagon.png", text)
    create_shape_image_with_text("arrow", "arrow.png", text)

    # Criar um documento Word no modo paisagem
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE

    # Ajustar a largura e altura da página no modo paisagem
    new_width, new_height = section.page_height, section.page_width
    section.page_width, section.page_height = new_width, new_height

    # Adicionar as imagens das formas geométricas ao documento
    doc.add_heading("Formas Geométricas", level=1)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # Centralizar o parágrafo

    run = paragraph.add_run()
    run.add_picture("circle.png", width=Inches(1))
    run.add_text("  ")  # Adicionar espaço entre as imagens
    run.add_picture("rectangle.png", width=Inches(1))
    run.add_text("  ")
    run.add_picture("diamond.png", width=Inches(1))
    run.add_text("  ")
    run.add_picture("hexagon.png", width=Inches(1))
    run.add_text("  ")
    run.add_picture("arrow.png", width=Inches(1))

    # Salvar o documento
    doc.save("formas_geometricas.docx")
    cmd.run("echo creation finished")


@task
def testreport2(cmd):
    """Report test 2"""
    cmd.run("echo creation started")

    # Função para desenhar as figuras geométricas com texto
    def draw_shapes(c: canvas.Canvas):
        _, height = landscape(letter)
        c.translate(inch, height - inch)  # Mover a origem (0,0) para a posição desejada
        c.setFont("Helvetica", 12)

        shapes = [
            ("circle", 0),
            ("rectangle", 2 * inch),
            ("diamond", 4 * inch),
            ("hexagon", 6 * inch),
        ]

        text = "Comodato teste 2024/06/12"

        for shape, x_offset in shapes:
            c.saveState()
            c.translate(x_offset, -3 * inch)
            text_width = c.stringWidth(text, "Helvetica", 12)
            if shape == "circle":
                c.setFillColorRGB(0.25, 0.91, 0.24)
                c.circle(1 * inch, -1 * inch, 1 * inch, stroke=1, fill=1)
                c.setFillColor("black")
                c.drawString((2 * inch - text_width) / 2, -1 * inch - 0.1 * inch, text)
            elif shape == "rectangle":
                c.setFillColorRGB(0.96, 0.63, 0.99)
                c.rect(0, -1.5 * inch, 2 * inch, inch, stroke=1, fill=1)
                c.setFillColor("black")
                c.drawString((2 * inch - text_width) / 2, -1 * inch - 0.1 * inch, text)
            elif shape == "diamond":
                # Desenhar e preencher um losango sem usar rotate
                diamond_points = [
                    (1 * inch, 0),
                    (2 * inch, -1 * inch),
                    (1 * inch, -2 * inch),
                    (0, -1 * inch),
                ]
                p = c.beginPath()
                p.moveTo(diamond_points[0][0], diamond_points[0][1])
                p.lineTo(diamond_points[1][0], diamond_points[1][1])
                p.lineTo(diamond_points[2][0], diamond_points[2][1])
                p.lineTo(diamond_points[3][0], diamond_points[3][1])
                p.close()
                c.setFillColorRGB(0.99, 0.91, 0.63)
                c.drawPath(p, stroke=1, fill=1)
                c.setFillColor("black")
                c.drawString((2 * inch - text_width) / 2, -1 * inch - 0.1 * inch, text)
            elif shape == "hexagon":
                # Desenhar e preencher um hexágono sem usar translate
                center_x = 1 * inch
                center_y = -1 * inch
                radius = 1 * inch
                hexagon_points = [
                    (
                        center_x + radius * math.cos(math.radians(angle)) * 1.5,
                        center_y + radius * math.sin(math.radians(angle)) / 1.5,
                    )
                    for angle in range(0, 360, 60)
                ]
                p = c.beginPath()
                p.moveTo(hexagon_points[0][0], hexagon_points[0][1])
                for point in hexagon_points[1:]:
                    p.lineTo(point[0], point[1])
                p.close()
                c.setFillColorRGB(0.69, 0.56, 0.44)
                c.drawPath(p, stroke=1, fill=1)
                c.setFillColor("black")
                c.drawString((2 * inch - text_width) / 2, -1 * inch - 0.1 * inch, text)

            # Desenhar uma seta formada por uma linha e um triângulo alinhados
            start_x = 1 * inch
            start_y = -1 * inch
            length = 1.5 * inch  # Comprimento total da seta
            width = 0.15 * inch  # Largura da haste da seta

            # Desenhar a haste da seta (linha)
            line_length = length * 0.2
            c.setLineWidth(2)
            c.line(start_x, start_y, start_x + line_length, start_y)

            # Desenhar a ponta da seta (triângulo preenchido)
            triangle_base = width
            triangle_height = length * 0.2
            arrow_tip = [
                (start_x + line_length, start_y - triangle_base),
                (start_x + line_length + triangle_height, start_y),
                (start_x + line_length, start_y + triangle_base),
            ]

            # Usar path para criar e preencher o triângulo
            c.setFillColorRGB(0, 0, 0)
            c.setStrokeColorRGB(0, 0, 0)
            p = c.beginPath()
            p.moveTo(arrow_tip[0][0], arrow_tip[0][1])
            p.lineTo(arrow_tip[1][0], arrow_tip[1][1])
            p.lineTo(arrow_tip[2][0], arrow_tip[2][1])
            p.close()
            c.drawPath(p, stroke=1, fill=1)
            c.restoreState()

    with Session_db() as session:
        ReportService().report_asset_timeline(3382, session)

    cmd.run("echo creation finished")
